import os
from io import BytesIO
from datetime import date
from flask import current_app
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def format_fio(full_name):
    parts = full_name.split()
    last_name = parts[0]
    if len(parts) == 3:
        initials = f"{parts[1][0]}.{parts[2][0]}."
        return f"{last_name} {initials}"
    elif len(parts) == 2:
        initials = f"{parts[1][0]}."
        return f"{last_name} {initials}"
    return full_name

def translate_date_to_ru(date: str) -> str:
    date = date.replace("May", "мая").replace("June", "июня").replace("July", "июля").replace("August", "августа")
    return date


def replace_placeholder(doc, placeholder, replacement):

    def replace_in_runs(runs):
        for run in runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement)

    # В параграфах
    for p in doc.paragraphs:
        if placeholder in p.text:
            replace_in_runs(p.runs)

    # В таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    for paragraph in cell.paragraphs:
                        replace_in_runs(paragraph.runs)


def sort_commission_members(members):
    """Сортирует членов комиссии с учетом всех вариантов написания ролей"""
    priority_order = [
        'председатель комиссии',
        'заместитель председателя',
        'член комиссии'
    ]

    def get_sort_key(member):
        formatted_role = format_role(member.role).lower()

        # Находим индекс роли в приоритетном списке
        for i, role in enumerate(priority_order):
            if role in formatted_role:
                return (i, member.user.full_name.lower())
        return (len(priority_order), member.user.full_name.lower())

    return sorted(members, key=get_sort_key)

def format_role(role):
    # форматирование роли
    role_lower = role.lower()
    if 'член' in role_lower:
        return 'член комиссии'
    elif any(x in role_lower for x in ['зам', 'заместитель']):
        return 'заместитель председателя'
    elif any(x in role_lower for x in ['председатель', 'пред']):
        return 'председатель комиссии'
    return role.capitalize()


def create_table_element(width="5000", border_size="4", border_color="000000"):
    tbl = OxmlElement('w:tbl')

    # Настройки таблицы
    tblPr = OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), width)
    tblW.set(qn('w:type'), "pct")
    tblPr.append(tblW)

    # Границы таблицы
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), border_size)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), border_color)
        tblBorders.append(border)
    tblPr.append(tblBorders)
    tbl.append(tblPr)

    return tbl


def create_cell(text, font='Arial', size="22", bold=False):
    # создаем ячейки, чтобы данные туда добавить
    tc = OxmlElement('w:tc')
    tcPr = OxmlElement('w:tcPr')
    tc.append(tcPr)

    p_elem = OxmlElement('w:p')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rPr.append(rFonts)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), size)
    rPr.append(sz)

    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = str(text)
    r.append(t)
    p_elem.append(r)
    tc.append(p_elem)

    return tc


def insert_table_after_paragraph(doc, placeholder, table_data, column_widths):
    # удаляем плейсхолдер
    for p in doc.paragraphs:
        if placeholder in p.text:
            parent = p._p.getparent()
            index = parent.index(p._p) + 1

            tbl = create_table_element()

            # Настройка ширины колонок
            tblGrid = OxmlElement('w:tblGrid')
            for width in column_widths:
                col = OxmlElement('w:gridCol')
                col.set(qn('w:w'), str(width))
                tblGrid.append(col)
            tbl.append(tblGrid)

            # Добавление данных
            for row in table_data:
                tr = OxmlElement('w:tr')
                for cell_text, cell_params in row:
                    tc = create_cell(cell_text, **cell_params)
                    tr.append(tc)
                tbl.append(tr)

            parent.insert(index, tbl)
            p.text = p.text.replace(placeholder, '')
            break


def generate_protocol(applicant, exam_date):
    tpl = os.path.join(current_app.config['TEMPLATES_FOLDER'], 'protocol_template.docx')
    doc = Document(tpl)

    # 1. Поиск председателя
    chairman = next(
        (m for m in exam_date.commission
         if any(x in m.role.lower() for x in ['председатель', 'пред'])),
        exam_date.commission[0] if exam_date.commission else None
    )

    replacements = {
        '{chairman}': chairman.user.full_name if chairman else '',
        '{date}': translate_date_to_ru(exam_date.date.strftime('«%d» %B %Yг.')),
        '{program_code}': exam_date.program.code,
        '{program_name}': exam_date.program.name,
        '{applicant_name}': applicant.full_name
    }

    for placeholder, value in replacements.items():
        replace_placeholder(doc, placeholder, value)

        # 2. Таблица комиссии
        sorted_commission = sort_commission_members(exam_date.commission)
        commission_data = [
            [['ФИО', {'bold': True}], ['Должность', {'bold': True}]]
        ]
        commission_data.extend(
            [[member.user.full_name, {}],
             [format_role(member.role) if 'член' not in format_role(member.role).lower() else '', {}]]
            for member in sorted_commission
        )
        insert_table_after_paragraph(
            doc,
            '{commission_table}',
            commission_data,
            [3500, 1500]
        )

    # 3. Таблица вопросов
    if not applicant.scores:
        # Если нет вопросов - добавляем отметку о неявке
        questions_data = [
            [['Примечание', {'bold': True, 'size': '24'}]],
            [['Неявка', {'size': '24'}]]
        ]
        insert_table_after_paragraph(
            doc,
            '{questions_table}',
            questions_data,
            [5000]  # Ширина колонки
        )

        # Удаляем плейсхолдер общего балла
        replace_placeholder(doc, '{total_score}', '0')
    else:
        total = sum(score.score for score in applicant.scores)
        questions_data = [
            [['№', {'bold': True}], ['Вопрос', {'bold': True}], ['Балл', {'bold': True}]]
        ]
        questions_data.extend(
            [[str(i), {}], [score.question.text, {}], [str(score.score), {}]]
            for i, score in enumerate(applicant.scores, start=1)
        )
        questions_data.append([['ИТОГО', {'bold': True}], ['', {}], [str(total), {'bold': True}]])

        insert_table_after_paragraph(
            doc,
            '{questions_table}',
            questions_data,
            [100, 4400, 500]  # Ширина колонок
        )

    # 4. Замена общего балла
    for p in doc.paragraphs:
        if '{total_score}' in p.text:
            p.text = p.text.replace('{total_score}', str(total))
            break

    # 5. Подписи комиссии
    signatures_data = [
        [['ФИО', {'bold': True}], ['Подпись', {'bold': True}]]
    ]
    signatures_data.extend(
        [[member.user.full_name, {}], ['', {}]]
        for member in sorted_commission
    )
    insert_table_after_paragraph(
        doc,
        '{signatures}',
        signatures_data,
        [3500, 1500]
    )

    return doc


def generate_vedomost(exam_date):
    template_path = os.path.join(current_app.config['TEMPLATES_FOLDER'], 'vedomost_template.docx')
    doc = Document(template_path)

    # 1. Поиск председателя
    chairman = next(
        (m for m in exam_date.commission
         if any(x in m.role.lower() for x in ['председатель', 'пред'])),
        exam_date.commission[0] if exam_date.commission else None
    )

    # 2. Подготовка данных для замены
    replacements = {
        '{date}': translate_date_to_ru(exam_date.date.strftime('«%d» %B %Yг.')),
        '{program_code}': exam_date.program.code,
        '{program_name}': f"{exam_date.program.name} / {exam_date.program.oop[0].name}" if exam_date.program.oop else exam_date.program.name,
        '{chairman}': format_fio(chairman.user.full_name) if chairman else ""
    }

    # 3. Замена плейсхолдеров в документе
    for placeholder, value in replacements.items():
        replace_placeholder(doc, placeholder, value)

    # 4. Таблица абитуриентов
    questions_data = [
        [['№ п/п', {'bold': True, 'size': '24'}], ['ФИО поступающего', {'bold': True, 'size': '24'}], ['Балл за ВИ', {'bold': True, 'size': '24'}]]
    ]
    # Фильтруем абитуриентов - оставляем только тех, у кого есть вопросы (даже если все ответы 0)
    present_applicants = [a for a in exam_date.applicants if getattr(a, 'scores', None)]
    sorted_applicants = sorted(present_applicants, key=lambda x: x.full_name.lower(), reverse=True)

    for i, applicant in enumerate(sorted_applicants, start=1):
        total_score = sum(score.score for score in applicant.scores) if applicant.scores else 0
        questions_data.append([
            [str(i), {'size': '24'}],
            [applicant.full_name, {'size': '24'}],
            [str(total_score), {'size': '24'}]
        ])

    insert_table_after_paragraph(
        doc,
        '{vedomost_table}',
        questions_data,
        [150, 4450, 400]  # Ширина колонок
    )

    sorted_commission = sort_commission_members(exam_date.commission)

    members_text = ""
    for member in sorted_commission:
        if member.role == 'председатель':
            continue

        members_text += " " * 65 + f" ____________________ / {format_fio(member.user.full_name)} /\n"
        members_text += " " * 65 + " ".ljust(12)+"(подпись)".ljust(30) + "(ФИО)\n\n"

    # Добавляем параграф с подписями
    p = doc.add_paragraph(members_text)

    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)

    return doc